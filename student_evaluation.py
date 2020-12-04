import tkinter as tk
from tkinter import ttk
from tkinter import Menu
from tkinter import messagebox
from tkinter import filedialog
from tkinter import font
from tkinter import simpledialog
from tkinter.simpledialog import _QueryString
import json
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.dml.color import ColorFormat
import pprint

categories = ["Content structure / ideas", "Language and Delivery", "Technical"]
colours = {
    4: "#F2F2F2",
    3: "#E6E6E6",
    2: "#D9D9D9",
    1: "#CCCCCC",
    0: "#BFBFBF",
}

evaluation_config = [
    [
        {
            "label": "Focus",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Purpose of presentation is clear from the outset. Supporting ideas maintain clear focus on the topic.",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Topic of the presentation is clear. Content generally supports the purpose",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Presentation lacks clear direction. Big ideas not specifically identified.",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "No focus at all. Audience cannot determine purpose of presentation.",
                },
            ],
        },
        {
            "label": "Organization",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Student presents information in logical, interesting sequence that audience follows",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Student presents information in logical sequence that audience can follow",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Audience has difficulty following because student jumps around.",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Audience cannot understand because there is no sequence of information",
                },
            ],
        },
        {
            "label": "Visual Aids",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Visual aids are readable, clear and professional looking, enhancing the message.",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Visual aids are mostly readable, clear and professional looking.",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Significant problems with readability, clarity, professionalism of visual aids.",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Visual aids are all unreadable, unclear and/or unprofessional.",
                },
            ],
        },
        {
            "label": "Question & Answer",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Speaker has prepared relevant questions for opening up the discussion and is able to stimulate discussion.",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Speaker has prepared relevant questions for opening up the discussion and is somewhat able to stimulate discussion",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Speaker has prepared questions but is not really able to stimulate discussion",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Speaker has not prepared questions",
                },
            ],
        },
    ],
    [
        {
            "label": "Eye Contact",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Holds attention of entire audience with the use of direct eye contact, seldom looking at notes",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Consistent use of direct eye contact with audience, but often returns to notes",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Displays minimal eye contact with audience, while reading mostly from the notes.",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "No eye contact with audience; entire presentation is read from notes",
                },
            ],
        },
        {
            "label": "Enthusiasm",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Demonstrates a strong, positive feeling about topic during entire presentation.",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Mostly shows positive feelings about topic.",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Shows some negativity toward topic presented.",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Shows no interest in topic presented.",
                },
            ],
        },
        {
            "label": "Elocution",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Student uses a clear voice so that all audience members can hear presentation",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Student’s voice is clear. Most audience members can hear presentation.",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Student’s voice is low. Audience has difficulty hearing presentation.",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Student mumbles, speaks too quietly for a majority of audience to hear.",
                },
            ],
        },
    ],
    [
        {
            "label": "Knowledge",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Demonstrate clear knowledge and understanding of the subject",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Show clear knowledge and understanding of most subject area",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Show some knowledge and understanding of the subject area",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Show no knowledge and understanding of the subject area",
                },
            ],
        },
        {
            "label": "Research",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Evidence of thorough research and preparation",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Evidence of sufficient research and preparation",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Evidence of some research and preparation",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Evidence of no research and preparation",
                },
            ],
        },
        {
            "label": "Discussion of new ideas",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Demonstrate  thorough knowledge while discussing new ideas",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Show  sufficient knowledge while discussing new ideas",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Show  some knowledge while discussing new ideas",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Show  no knowledge while discussing new ideas",
                },
            ],
        },
        {
            "label": "Argument",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Opinion set out in a concise and persuasive manner",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Opinion is not concise and persuasive manner",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Opinion is clearly demonstrated but not persuasive ",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Opinion is not demonstrated or highlighted",
                },
            ],
        },
        {
            "label": "Questions",
            "evaluations": [
                {
                    "label": "Excellent",
                    "value": 4,
                    "description": "Responded very well to technical questions",
                },
                {
                    "label": "Good",
                    "value": 3,
                    "description": "Could answer most technical questions related to the presentation",
                },
                {
                    "label": "Fair",
                    "value": 2,
                    "description": "Could answer some technical questions related to the presentation",
                },
                {
                    "label": "Poor",
                    "value": 1,
                    "description": "Could not answer any technical questions related to the presentation",
                },
            ],
        },
    ],
]

evaluation_radios = []
evaluation_descriptions = []
evaluation_choices = []
evaluation_comments = []
total_marks = 0
all_students = []
student_menus = []
current_student_index: int = 0

default_student = {
    "student_first_name": "New",
    "student_last_name": "Student",
    "student_id": "00000000",
    "student_symbol": "",
    "student_marks": 0,
    "student_comment": "",
    "evaluations": [],
    "result_comment": ""
}


def save_student_button_push():
    global current_student_index
    all_students[current_student_index] = get_student()
    update_student_menus()
    update_select_student_menu()


def menu_select_student(menu_index: int):
    def x():
        all_students[current_student_index] = get_student()
        select_student(menu_index)
        current_student.set(student_menus[menu_index])

    return x


def select_student(menu_index: int):
    global current_student_index
    current_student_index = menu_index
    set_student(all_students[current_student_index])


def add_new_student():
    all_students.append(default_student)
    select_student(len(all_students) - 1)


def get_student():
    student = {
        "student_first_name": student_first_name.get(),
        "student_last_name": student_last_name.get(),
        "student_id": student_id.get(),
        "student_symbol": student_symbol.get(),
        "student_marks": get_marks(),
        "student_comment": student_comment.get(),
        "evaluations": [],
        "result_comment": result_comment.get(),
    }
    txt = "{student_first_name:s} {student_last_name:s}"
    student["student_name"] = txt.format(student_first_name=student["student_first_name"],
                                         student_last_name=student["student_last_name"])
    #  comment
    for i in range(len(evaluation_comments)):
        student["evaluations"].append({
            "choice": evaluation_choices[i].get(),
            "comment": evaluation_comments[i].get()
        })

    return student


def set_student(student):
    student_first_name.set(student["student_first_name"])
    student_last_name.set(student["student_last_name"])
    student_id.set(student["student_id"])
    student_symbol.set(student["student_symbol"])
    #  comment
    student_comment.set(student["student_comment"])
    for i in range(len(evaluation_comments)):
        evaluation_choices[i].set(student["evaluations"][i]["choice"])
        evaluation_comments[i].set(student["evaluations"][i]["comment"])
    result_comment.set(student["result_comment"])


def update_student_menus():
    student_menus.clear()
    for i in range(len(all_students)):
        student = all_students[i]
        txt = "{i:d} : {first_name:s} {last_name:s} ({student_id:s})"
        label = txt.format(i=i,
                           first_name=student["student_first_name"],
                           last_name=student["student_last_name"],
                           student_id=student["student_id"])
        student_menus.append(label)


def add_student_button_push():
    global current_student_index
    all_students[current_student_index] = get_student()
    add_new_student()
    update_student_menus()
    update_select_student_menu()


def update_select_student_menu():
    global current_student_index
    menu = select_student_menu["menu"]
    menu.delete(0, "end")
    for i in range(len(student_menus)):
        menu.add_command(label=student_menus[i], command=menu_select_student(i))
    current_student.set(student_menus[current_student_index])


def get_result_string(result_marks):
    f = "Result: {result_marks:d} / {total_marks}"
    mark_string = f.format(result_marks=result_marks, total_marks=total_marks)
    return mark_string


def get_marks():
    mark_sum = 0
    y = 0
    for i in range(len(categories)):
        for j in range(len(evaluation_config[i])):
            evaluation_parameter = evaluation_config[i][j]
            evaluation_count = len(evaluation_parameter["evaluations"])

            selected = evaluation_choices[y].get()
            if selected < evaluation_count:
                mark_sum = mark_sum + evaluation_parameter["evaluations"][selected]["value"]

            y = y + 1

    return mark_sum


def click_evaluation(y, x):
    def x():
        result_label["text"] = get_result_string(get_marks())

    return x


"""
def click_evaluation_selection(clicked_i, clicked_j, clicked_y):
    def x():
        y = 0
        for i in range(len(categories)):
            for j in range(len(evaluation_config[i])):
                evaluation_parameter = evaluation_config[i][j]
                evaluation_count = len(evaluation_parameter["evaluations"])
                for k in range(len(evaluation_parameter["evaluations"]) + 1):
                    if y == clicked_y:
                        evaluation_descriptions[y][k].grid()
                y = y + 1

    return x
"""


def create_save_data_dict():
    data = {
        "topic": topic_entered.get(),
        "total_marks": total_marks,
        "students": all_students,
    }

    return data


def create_result_json():
    data = create_save_data_dict()
    return json.dumps(data)


"""
def validate_fill_in():
    message = ""
    result = True

    if student_name_first_entered.get() == "":
        message += "* Missing Student name (first name)\n"

    if student_name_last_entered.get() == "":
        message += "* Missing Student name (last name)\n"

    if student_id_entered.get() == "":
        message += "* Missing Student ID\n"

    if topic_entered.get() == "":
        message += "* Missing Topic of presentation\n"

    y = 0
    for i in range(len(categories)):
        for j in range(len(evaluation_config[i])):
            evaluation_parameter = evaluation_config[i][j]
            evaluation_count = len(evaluation_parameter["evaluations"])

            if evaluation_count < evaluation_choices[y].get():
                txt = "* {label:s} is not assessed.\n"
                message += txt.format(label=evaluation_parameter["label"])

            y = y + 1

    if message != "":
        messagebox.showinfo("Error", message)
        result = False

    return result
"""


def validate_fill_in():
    messages: str = ""
    result = True

    if topic_entered.get() == "":
        messages += "* Missing Topic of presentation\n"

    for i in range(len(student_menus)):
        student_label: str = student_menus[i]
        message: str = ""

        student = all_students[i]

        if student["student_first_name"] == "":
            message += "* Missing Student name (first name)\n"

        if student["student_last_name"] == "":
            message += "* Missing Student name (last name)\n"

        if student["student_id"] == "":
            message += "* Missing Student ID\n"

        if student["student_symbol"] == "":
            message += "* Missing Student Symbol\n"

        y = 0
        for i in range(len(categories)):
            for j in range(len(evaluation_config[i])):
                evaluation_parameter = evaluation_config[i][j]
                evaluation_count = len(evaluation_parameter["evaluations"])

                if evaluation_count < student["evaluations"][y]["choice"]:
                    txt = "* {label:s} is not assessed.\n"
                    message += txt.format(label=evaluation_parameter["label"])

                y = y + 1

        if message != "":
            if messages != "":
                messages += "\n"
            messages += "Student: " + student_label + "\n"
            messages += message

    if messages != "":
        messagebox.showinfo("Error", messages)
        result = False

    return result


def menu_save():
    if validate_fill_in():
        json_string = create_result_json()

        txt = "{student_id:s}.json"
        file_name = txt.format(txt, student_id=student_id_entered.get())

        with open(file_name, mode='w') as f:
            f.write(json_string)


def menu_export_as_word():
    all_students[current_student_index] = get_student()
    if validate_fill_in():
        txt = "{student_id:s}.docx"
        filename = txt.format(student_id=student_id_entered.get())
        path = tk.filedialog.asksaveasfile(mode="w", initialfile=filename, defaultextension="docx")
        filename = path.name
        path.close()
        do_export_as_word(filename)


def do_export_as_word(filename):
    save_data = create_save_data_dict()
    pprint.pprint(save_data)

    doc = Document()
    txt = "Topic of Presentation: {topic:s}"
    doc.add_paragraph(txt.format(topic=save_data["topic"]))

    student_count = len(save_data["students"])
    table = doc.add_table(rows=student_count, cols=5)
    table.style = 'TableNormal'

    #
    # Student Information
    #
    mark_txt = "{result_marks:d} / {total_marks:d}"
    student_symbols_by_evaluation_choice = []
    row_index = 0
    for category_idx in range(len(evaluation_config)):
        for category_row_index in range(len(evaluation_config[category_idx])):
            evaluations = evaluation_config[category_idx][category_row_index]["evaluations"]
            student_symbols_by_evaluation_choice.append([])
            for evaluation_index in range(len(evaluations)):
                student_symbols_by_evaluation_choice[row_index].append([])
            row_index += 1

    for i in (range(student_count)):
        student = save_data["students"][i]
        student_result = mark_txt.format(result_marks=student["student_marks"],
                                         total_marks=total_marks)
        txt = "{student_name:s} ({student_symbol:s})"
        student_name = txt.format(student_name=student["student_name"],
                                  student_symbol=student["student_symbol"])
        table.cell(i, 0).text = "Student name:"
        table.cell(i, 1).text = student_name
        table.cell(i, 2).text = "Student ID:"
        table.cell(i, 3).text = student["student_id"]
        table.cell(i, 4).text = student_result

        if student["student_comment"] != "":
            table.cell(i, 1).paragraphs[0].runs[0].add_comment(student["student_comment"], "Student Evaluation App")

        if student["result_comment"] != "":
            table.cell(i, 4).paragraphs[0].runs[0].add_comment(student["result_comment"], "Student Evaluation App")

        for j in range(len(student["evaluations"])):
            choice = student["evaluations"][j]["choice"]
            txt = "{j:d} {choice:d}"
            student_symbols_by_evaluation_choice[j][choice].append(student["student_symbol"])

    pprint.pprint(student_symbols_by_evaluation_choice)

    table = doc.add_table(rows=13, cols=5)
    table.style = 'TableGrid'

    cell_row_index = 0
    evaluate_index = 0

    for category_idx in range(len(evaluation_config)):
        table.cell(cell_row_index, category_idx).text = categories[category_idx]
        table.cell(cell_row_index, category_idx).paragraphs[0].runs[0].font.bold = True

        #
        # label
        #
        for category_row_index in range(len(evaluation_config[category_idx])):
            options = evaluation_config[category_idx][category_row_index]["evaluations"]

            for option_index in range(len(options)):
                option = options[option_index]
                txt = "{score:d} - {label:s}"
                table.cell(cell_row_index, option_index + 1).text = txt.format(score=option["value"], label=option["label"])
                table.cell(cell_row_index, option_index + 1).paragraphs[0].runs[0].font.bold = True

        for category_row_index in range(len(evaluation_config[category_idx])):
            choices = evaluation_config[category_idx][category_row_index]["evaluations"]
            cell_row_index += 1
            option = evaluation_config[category_idx][category_row_index]
            table.cell(cell_row_index, 0).text = option["label"]
            if evaluation_comments[evaluate_index].get() != "":
                table.cell(cell_row_index, 0).paragraphs[0].runs[0].add_comment(evaluation_comments[evaluate_index].get(),
                                                                         "Student Evaluation App")

            for choice_index in range(len(choices)):
                choice = choices[choice_index]
                table.cell(cell_row_index, choice_index + 1).text = choice["description"]

                student_symbols = ""
                txt = "({symbol:s})"
                for student_symbol in student_symbols_by_evaluation_choice[evaluate_index][choice_index]:
                    student_symbols += txt.format(symbol=student_symbol)
                if student_symbols != "":
                    table.cell(cell_row_index, choice_index + 1).text += "\n" + student_symbols

            evaluate_index += 1

    doc.save(filename)
    messagebox.showinfo("title", "save as " + filename)


def menu_exit():
    win.quit()
    win.destroy()
    exit()


def menu_about():
    messagebox.showinfo("About",
                        "This application is for student evaluation created by SUX2 group in SOFT808 in S3 2020.")


"""
class CommentDialog(simpledialog.Dialog):
    def __init__(self, master,
                 text='', buttons=[], default=None, cancel=None,
                 title=None, class_=None):
        pass

    def body(self, master):
        self.comment = tk.StringVar()
        self.comment_entered = ttk.Entry(self.comment, width=12, textvariable=self.comment)
        self.comment_entered.pack(side=tk.LEFT, padx=5, pady=5)

    # def buttonbox(self):
"""


class CommentDialog(_QueryString):
    def __init__(self, *args, **kw):
        _QueryString.__init__(self, *args, **kw)

    def body(self, master):
        entry = _QueryString.body(self, master)
        entry.configure(width=50)
        return entry

    def getresult(self):
        return self.entry.get()


def show_comment_dialog(label, comment):
    def x():
        prompt = "Enter Comment for {label:s}"
        comment.set(CommentDialog("Enter Comment.", prompt.format(label=label), initialvalue=comment.get()).result)

    return x


def reset_parameters():
    student_first_name.set("")
    student_last_name.set("")
    student_id.set("")
    topic.set("")
    for i in range(len(evaluation_choices)):
        evaluation_choices[i].set(4)

    result_label["text"] = get_result_string(get_marks())


def initialize_values():
    for i in range(len(categories)):
        evaluation_count = len(evaluation_config[i])
        for j in range((evaluation_count)):
            default_student["evaluations"].append({
                "choice": evaluation_count,
                "comment": ""
            })
            evaluation_choices.append(tk.IntVar())
            evaluation_comments.append(tk.StringVar())


def build_evaluation_form():
    global total_marks

    y = 0
    tab_control = ttk.Notebook(evaluation_frame)
    tab_control.grid(column=0, row=0)
    tab_frames = []
    for i in range(len(categories)):
        category_frame = ttk.LabelFrame(tab_control)
        tab_control.add(category_frame, text=categories[i])
        # category_frame.grid(column = 0, row = y * 2, sticky = tk.W)
        tab_frames.append(category_frame)

        for j in range(len(evaluation_config[i])):
            evaluation_radios.append([])
            evaluation_descriptions.append([])
            evaluation_parameter = evaluation_config[i][j]
            evaluation_count = len(evaluation_parameter["evaluations"])

            max_mark = 0
            for k in range(evaluation_count):
                colour = colours[evaluation_parameter["evaluations"][k]["value"]]
                txt = "{value:d} - {label:s}"
                r = tk.Radiobutton(category_frame,
                                   text=txt.format(value=evaluation_parameter["evaluations"][k]["value"],
                                                 label=evaluation_parameter["evaluations"][k]["label"]),
                                   variable=evaluation_choices[y], width=20, bg=colour,
                                   value=k, command=click_evaluation(y, k))
                r.grid(column=k + 2, row=y * 2, sticky=tk.W)

                description = evaluation_parameter["evaluations"][k]["description"]
                d = tk.Message(category_frame, text=description, bg=colour, font=small_font)
                d.grid(column=k + 2, row=y * 2 + 1, sticky=tk.W + tk.E + tk.N + tk.S)

                evaluation_radios[y].append(r)
                evaluation_descriptions[y].append(d)

                if max_mark < evaluation_parameter["evaluations"][k]["value"]:
                    max_mark = evaluation_parameter["evaluations"][k]["value"]

            total_marks = total_marks + max_mark

            label = "Not yet"
            r = tk.Radiobutton(category_frame, text=label, variable=evaluation_choices[y],
                               value=evaluation_count, bg=colours[0], command=click_evaluation(y, evaluation_count))
            r.grid(column=evaluation_count + 2, row=y * 2, sticky=tk.W)

            d = tk.Message(category_frame, text="", bg=colours[0], font=small_font)
            d.grid(column=evaluation_count + 2, row=y * 2 + 1, sticky=tk.W + tk.E + tk.N + tk.S)
            # d.grid_remove()

            evaluation_radios[y].append(r)
            evaluation_descriptions[y].append(d)

            category_label = ttk.Label(category_frame, text=evaluation_parameter["label"], width=20)
            category_label.grid(column=1, row=y * 2, sticky=tk.W)

            evaluation_comment = evaluation_comments[y]
            evaluation_comment_button = ttk.Button(category_frame,
                                                   text="+Comment",
                                                   command=show_comment_dialog(evaluation_parameter["label"],
                                                                               evaluation_comment))
            evaluation_comment_button.grid(column=1, row=y * 2 + 1, sticky=tk.W + tk.N)
            evaluation_comments[y] = evaluation_comment

            y = y + 1

    result_label["text"] = get_result_string(0)


# Create instance
win = tk.Tk()
win.resizable(0, 0)

# Add a title
win.title("Student Evaluation")

default_font = tk.font.nametofont("TkDefaultFont")
small_font = tk.font.Font(size=default_font.cget("size") - 1)

#
#  Frames
#

select_student_frame = ttk.LabelFrame(win, text="Select Student")
select_student_frame.grid(column=0, row=0, columnspan=2, padx=8, pady=4, sticky=tk.W)

student_frame = ttk.LabelFrame(win, text="Student")
student_frame.grid(column=0, row=1, padx=8, pady=4, sticky=tk.W)

result_frame = ttk.LabelFrame(win, text="Result")
result_frame.grid(column=1, row=1, padx=8, pady=4, sticky=tk.E + tk.W + tk.N + tk.S)

evaluation_frame = ttk.LabelFrame(win, text="Evaluation")
evaluation_frame.grid(column=0, row=2, columnspan=2, padx=8, pady=4)

#
#  Student Menu
#
current_student = tk.StringVar()
current_student.set("")
select_student_menu = tk.OptionMenu(select_student_frame, current_student, "")
select_student_menu.grid(column=0, row=0, sticky=tk.W)

save_student_button = ttk.Button(select_student_frame, text="Save Student", command=save_student_button_push)
save_student_button.grid(column=1, row=0, sticky=tk.W)
add_student_button = ttk.Button(select_student_frame, text="Add New Student", command=add_student_button_push)
add_student_button.grid(column=1, row=1, sticky=tk.W)

#
#  Topic of presentation
#
label = ttk.Label(select_student_frame, text="Topic of presentation")
label.grid(column=2, row=0, sticky=tk.W, ipady=4)

topic = tk.StringVar()
topic_entered = ttk.Entry(select_student_frame, width=50, textvariable=topic)
topic_entered.grid(column=3, row=0, columnspan=2, sticky=tk.W)

student_comment = tk.StringVar()
student_frame_comment_button = ttk.Button(student_frame, text="+Comment",
                                          command=show_comment_dialog("Student", student_comment))
student_frame_comment_button.grid(column=0, row=3, sticky=tk.W)

#
#  Student name
#
label = ttk.Label(student_frame, text="Student name")
label.grid(column=0, row=1, sticky=tk.W, ipady=4)

label = ttk.Label(student_frame, text="First name", font=small_font)
label.grid(column=1, row=0, sticky=tk.W)
label = ttk.Label(student_frame, text="Last name", font=small_font)
label.grid(column=2, row=0, sticky=tk.W)

student_first_name = tk.StringVar()
student_last_name = tk.StringVar()
student_name_first_entered = ttk.Entry(student_frame, width=15, textvariable=student_first_name)
student_name_first_entered.grid(column=1, row=1, sticky=tk.W)
student_name_last_entered = ttk.Entry(student_frame, width=15, textvariable=student_last_name)
student_name_last_entered.grid(column=2, row=1, sticky=tk.W)

#
#  Student ID
#
label = ttk.Label(student_frame, text="Student ID")
label.grid(column=5, row=1, sticky=tk.W, ipadx=10, ipady=4)

student_id = tk.StringVar()
student_id_entered = ttk.Entry(student_frame, width=12, textvariable=student_id)
student_id_entered.grid(column=6, row=1, sticky=tk.W)

#
#  Student Symbol
#
label = ttk.Label(student_frame, text="Student Symbol")
label.grid(column=0, row=2, sticky=tk.W, ipadx=10, ipady=4)

student_symbol = tk.StringVar()
student_symbol_entered = ttk.Entry(student_frame, width=3, textvariable=student_symbol)
student_symbol_entered.grid(column=1, row=2, sticky=tk.W)

#
#  Result
#
result_label = ttk.Label(result_frame, text=get_result_string(0), width=30)
result_label.grid(column=0, row=0, sticky=tk.E)

export_button = ttk.Button(result_frame, text="Export as Word file", command=menu_export_as_word, width=20)
export_button.grid(column=1, row=0, sticky=tk.E + tk.N + tk.S)

result_comment = tk.StringVar()
result_frame_comment_button = ttk.Button(result_frame, text="+Comment",
                                         command=show_comment_dialog("Result", result_comment))
result_frame_comment_button.grid(column=0, row=3, sticky=tk.W + tk.S)

result_reset_button = ttk.Button(result_frame, text="Reset", command=reset_parameters, width=20)
result_reset_button.grid(column=1, row=1, sticky=tk.E + tk.N + tk.S, pady=10)

initialize_values()
build_evaluation_form()

menu_bar = Menu(win)
win.config(menu=menu_bar)

file_menu = Menu(menu_bar, tearoff=0)
# file_menu.add_command(label = "Save", command = menu_save)
file_menu.add_command(label="Export as Word", command=menu_export_as_word)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=menu_exit)
menu_bar.add_cascade(label="File", menu=file_menu)

help_menu = Menu(menu_bar, tearoff=0)
help_menu.add_command(label="About", command=menu_about)
menu_bar.add_cascade(label="Help", menu=help_menu)

add_new_student()
select_student(0)
update_student_menus()
update_select_student_menu()

win.mainloop()
